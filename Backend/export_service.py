"""
Export Service for Boise Prosthodontics AI Scribe
Provides functionality to export data in various formats (PDF, DOCX, CSV, ZIP)
"""

import io
import json
import csv
import zipfile
from pathlib import Path
from datetime import datetime
from typing import Optional
import logging

# PDF Generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# DOCX Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import get_session_by_id, get_sessions_by_provider, get_provider_by_id


class ExportService:
    """Service for exporting data in various formats"""
    
    def __init__(self):
        """Initialize export service"""
        self.voice_profiles_dir = Path("/app/voice_profiles")
        logging.info("✅ Export service initialized")
    
    def export_session_to_pdf(self, session_id: str) -> bytes:
        """
        Generate PDF with session transcript and SOAP note
        
        Args:
            session_id: Session identifier
            
        Returns:
            bytes: PDF file content
        """
        try:
            # Get session data
            session = get_session_by_id(session_id)
            if not session:
                raise ValueError(f"Session {session_id} not found")
            
            # Create PDF buffer
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#3B82F6'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#1F2937'),
                spaceAfter=12,
                spaceBefore=12
            )
            normal_style = styles['BodyText']
            
            # Title
            title = Paragraph("Medical Transcription & SOAP Note", title_style)
            elements.append(title)
            elements.append(Spacer(1, 12))
            
            # Session Information Table
            session_info = [
                ['Session ID:', session.get('session_id', 'N/A')],
                ['Date:', session.get('timestamp', datetime.now()).strftime('%Y-%m-%d %H:%M')],
                ['Provider:', session.get('doctor_name', 'N/A')],
                ['Patient:', session.get('patient_name', 'N/A')],
                ['Patient ID:', session.get('patient_id', 'N/A')],
            ]
            
            info_table = Table(session_info, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F3F4F6')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 20))
            
            # Transcript Section
            transcript_heading = Paragraph("Transcript", heading_style)
            elements.append(transcript_heading)
            
            transcript_text = session.get('transcript', 'No transcript available')
            transcript_para = Paragraph(transcript_text.replace('\n', '<br/>'), normal_style)
            elements.append(transcript_para)
            elements.append(Spacer(1, 20))
            
            # SOAP Note Section
            soap_heading = Paragraph("SOAP Note", heading_style)
            elements.append(soap_heading)
            
            soap_text = session.get('soap_note', 'No SOAP note available')
            
            # Parse SOAP note sections if in JSON format
            try:
                soap_data = json.loads(soap_text)
                for section, content in soap_data.items():
                    section_title = Paragraph(f"<b>{section.upper()}</b>", normal_style)
                    elements.append(section_title)
                    section_content = Paragraph(str(content).replace('\n', '<br/>'), normal_style)
                    elements.append(section_content)
                    elements.append(Spacer(1, 10))
            except (json.JSONDecodeError, AttributeError):
                # Plain text SOAP note
                soap_para = Paragraph(soap_text.replace('\n', '<br/>'), normal_style)
                elements.append(soap_para)
            
            # Build PDF
            doc.build(elements)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            logging.info(f"✅ Generated PDF for session {session_id}")
            return pdf_bytes
            
        except Exception as e:
            logging.error(f"Error generating PDF for session {session_id}: {e}")
            raise
    
    def export_session_to_docx(self, session_id: str) -> bytes:
        """
        Generate Word document with session data
        
        Args:
            session_id: Session identifier
            
        Returns:
            bytes: DOCX file content
        """
        try:
            # Get session data
            session = get_session_by_id(session_id)
            if not session:
                raise ValueError(f"Session {session_id} not found")
            
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            title = doc.add_heading('Medical Transcription & SOAP Note', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(59, 130, 246)
            
            # Session Information
            doc.add_paragraph()
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            info_data = [
                ('Session ID:', session.get('session_id', 'N/A')),
                ('Date:', session.get('timestamp', datetime.now()).strftime('%Y-%m-%d %H:%M')),
                ('Provider:', session.get('doctor_name', 'N/A')),
                ('Patient:', session.get('patient_name', 'N/A')),
                ('Patient ID:', session.get('patient_id', 'N/A')),
            ]
            
            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Transcript Section
            doc.add_paragraph()
            doc.add_heading('Transcript', 1)
            transcript_text = session.get('transcript', 'No transcript available')
            doc.add_paragraph(transcript_text)
            
            # SOAP Note Section
            doc.add_paragraph()
            doc.add_heading('SOAP Note', 1)
            
            soap_text = session.get('soap_note', 'No SOAP note available')
            
            # Parse SOAP note sections if in JSON format
            try:
                soap_data = json.loads(soap_text)
                for section, content in soap_data.items():
                    section_heading = doc.add_heading(section.upper(), 2)
                    doc.add_paragraph(str(content))
            except (json.JSONDecodeError, AttributeError):
                # Plain text SOAP note
                doc.add_paragraph(soap_text)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            logging.info(f"✅ Generated DOCX for session {session_id}")
            return docx_bytes
            
        except Exception as e:
            logging.error(f"Error generating DOCX for session {session_id}: {e}")
            raise
    
    def export_sessions_to_csv(self, provider_id: Optional[str] = None, 
                               start_date: Optional[datetime] = None,
                               end_date: Optional[datetime] = None) -> str:
        """
        Export multiple sessions to CSV
        
        Args:
            provider_id: Filter by provider (optional)
            start_date: Start date filter (optional)
            end_date: End date filter (optional)
            
        Returns:
            str: CSV content
        """
        try:
            # Get sessions
            if provider_id:
                sessions = get_sessions_by_provider(int(provider_id))
            else:
                from database import get_all_sessions
                sessions = get_all_sessions()
            
            # Filter by date if provided
            if start_date or end_date:
                filtered_sessions = []
                for session in sessions:
                    session_date = session.get('timestamp')
                    if isinstance(session_date, str):
                        session_date = datetime.fromisoformat(session_date.replace('Z', '+00:00'))
                    
                    if start_date and session_date < start_date:
                        continue
                    if end_date and session_date > end_date:
                        continue
                    filtered_sessions.append(session)
                sessions = filtered_sessions
            
            # Create CSV
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Header
            writer.writerow([
                'Date',
                'Session ID',
                'Provider',
                'Patient Name',
                'Patient ID',
                'Template Used',
                'Sent to Dentrix',
                'Email Sent',
                'Dentrix Note ID'
            ])
            
            # Data rows
            for session in sessions:
                timestamp = session.get('timestamp', '')
                if isinstance(timestamp, datetime):
                    timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
                
                writer.writerow([
                    timestamp,
                    session.get('session_id', ''),
                    session.get('doctor_name', ''),
                    session.get('patient_name', ''),
                    session.get('patient_id', ''),
                    session.get('template_used', ''),
                    'Yes' if session.get('sent_to_dentrix') else 'No',
                    'Yes' if session.get('email_sent') else 'No',
                    session.get('dentrix_note_id', '')
                ])
            
            csv_content = output.getvalue()
            output.close()
            
            logging.info(f"✅ Generated CSV with {len(sessions)} sessions")
            return csv_content
            
        except Exception as e:
            logging.error(f"Error generating CSV: {e}")
            raise
    
    def export_voice_profile(self, provider_name: str) -> bytes:
        """
        Export voice profile as ZIP file
        
        Args:
            provider_name: Provider name
            
        Returns:
            bytes: ZIP file content
        """
        try:
            # Sanitize provider name for directory
            safe_name = provider_name.lower().replace(' ', '_')
            profile_dir = self.voice_profiles_dir / safe_name
            
            if not profile_dir.exists():
                raise ValueError(f"Voice profile not found for {provider_name}")
            
            # Create ZIP buffer
            buffer = io.BytesIO()
            
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Add profile.pkl if exists
                profile_file = profile_dir / 'profile.pkl'
                if profile_file.exists():
                    zip_file.write(profile_file, 'profile.pkl')
                
                # Add metadata.json if exists
                metadata_file = profile_dir / 'metadata.json'
                if metadata_file.exists():
                    zip_file.write(metadata_file, 'metadata.json')
                else:
                    # Create basic metadata
                    metadata = {
                        'provider_name': provider_name,
                        'exported_at': datetime.now().isoformat(),
                        'version': '1.0'
                    }
                    zip_file.writestr('metadata.json', json.dumps(metadata, indent=2))
                
                # Add any sample audio files
                for audio_file in profile_dir.glob('*.wav'):
                    zip_file.write(audio_file, f'samples/{audio_file.name}')
            
            zip_bytes = buffer.getvalue()
            buffer.close()
            
            logging.info(f"✅ Exported voice profile for {provider_name}")
            return zip_bytes
            
        except Exception as e:
            logging.error(f"Error exporting voice profile for {provider_name}: {e}")
            raise


# Global export service instance
export_service = ExportService()
